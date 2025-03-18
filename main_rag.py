import os
import re
import pinecone
from typing import List, Dict, Any, Optional
from dotenv import load_dotenv
import google.generativeai as genai
from sentence_transformers import SentenceTransformer
from pinecone.grpc import PineconeGRPC
from pinecone import ServerlessSpec
from docx import Document  # For handling .docx files


# ------------------------------------------------------------------------------
# Configuration
# ------------------------------------------------------------------------------
load_dotenv()

INDEX_NAME = "vendai"
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
PINECONE_ENVIRONMENT = os.getenv("PINECONE_ENVIRONMENT")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")


genai.configure(api_key=GEMINI_API_KEY)  # Configure Gemini API
model = genai.GenerativeModel('gemini-1.5-pro')  # Use gemini-pro model

# Checking for the Gemini models available to me
# for model in genai.list_models():
#     print(model)

if PINECONE_API_KEY:
    pc = PineconeGRPC(api_key=PINECONE_API_KEY)
else:
    pc = None
    print("Pinecone API key not found in environment variables.")

class RAGSystem:
    def __init__(
        self,
        pc_client: str,
        index_name: str,
        embedding_model_name: str,
        gemini_api_key: Optional[str] = None,
        dimension: int = 1024
    ):
        """
        Initialized the RAG system with Pinecone as the vector database and Gemini for generation.
        """
        if pc:
            existing_indexes = pc_client.list_indexes().names()
            if index_name not in existing_indexes:
                print(f"Creating index: {index_name}")
                pc_client.create_index(
                    name=index_name,
                    dimension=dimension,
                    metric="dotproduct",
                    spec=ServerlessSpec(cloud="aws", region="us-east-1")
                )
            self.index = pc_client.Index(index_name)
        else:
            self.index = None
            print("Pinecone connection not established. Check API key.")

        self.embedding_model = SentenceTransformer(embedding_model_name)
        self.dimension = dimension

        self.gemini_client = None
        if gemini_api_key:
            genai.configure(api_key=gemini_api_key)
            self.gemini_client = genai.GenerativeModel('gemini-2.0-pro-exp-02-05')

    def embed_text(self, text: str) -> List[float]:
        embedding = self.embedding_model.encode(text)
        return embedding.tolist()

    def load_and_split_document(self, file_path: str, chunk_delimiter: str = "########") -> List[str]:
        file_extension = os.path.splitext(file_path)[1].lower()
        text = ""

        try:
            if file_extension == ".txt":
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif file_extension == ".docx":
                document = Document(file_path)
                for paragraph in document.paragraphs:
                    text += paragraph.text + "\n\n"
                text = text.strip()
            else:
                raise ValueError(f"Unsupported file format: {file_extension}. Only .txt and .docx are supported.")
        except FileNotFoundError:
            print(f"Error: File not found at {file_path}")
            return []
        except Exception as e:
            print(f"An error occurred while loading the document: {e}")
            return []

        chunks = text.split(chunk_delimiter)
        return [chunk.strip() for chunk in chunks if chunk.strip()]

    def chunk_by_sentences(self, text: str, chunk_size: int = 3, chunk_overlap: int = 1) -> List[str]:
        sentences = self.split_text_into_sentences(text)
        chunks = []
        for i in range(0, len(sentences), chunk_size - chunk_overlap):
            chunk = " ".join(sentences[i : i + chunk_size])
            chunks.append(chunk)
        return chunks

    def add_documents(self, documents: List[Dict[str, Any]]):
        if not self.index:
            print("Pinecone index not initialized. Cannot add documents.")
            return

        vectors_to_upsert = []

        for doc in documents:
            doc_id = doc['id']
            text = doc['text']
            metadata = doc.get('metadata', {})

            if 'text' not in metadata:
                metadata['text'] = text

            embedding = self.embed_text(text)

            vectors_to_upsert.append({
                'id': str(doc_id),
                'values': embedding,
                'metadata': metadata
            })

            if len(vectors_to_upsert) >= 100:
                self.index.upsert(vectors=vectors_to_upsert)
                vectors_to_upsert = []

        if vectors_to_upsert:
            self.index.upsert(vectors=vectors_to_upsert)

    def index_document(self, file_path: str, chunk_delimiter: str = "\n\n", source: str = "file"):
        chunks = self.load_and_split_document(file_path, chunk_delimiter)
        documents_to_index = []
        for i, chunk in enumerate(chunks):
            doc_id = f"{source}-{i}"
            metadata = {"source": source}
            documents_to_index.append({"id": doc_id, "text": chunk, "metadata": metadata})
        self.add_documents(documents_to_index)
        print(f"Indexed {len(chunks)} chunks from {file_path}")

    def index_document_by_sentences(self, file_path: str, chunk_size: int = 3, chunk_overlap: int = 1, source: str = "file"):
        file_extension = os.path.splitext(file_path)[1].lower()
        full_text = ""

        try:
            if file_extension == ".txt":
                with open(file_path, 'r', encoding='utf-8') as f:
                    full_text = f.read()
            elif file_extension == ".docx":
                document = Document(file_path)
                for paragraph in document.paragraphs:
                    full_text += paragraph.text + "\n"
                full_text = full_text.strip()
            else:
                raise ValueError(f"Unsupported file format: {file_extension}. Only .txt and .docx are supported.")
        except FileNotFoundError:
            print(f"Error: File not found at {file_path}")
            return
        except Exception as e:
            print(f"An error occurred while loading the document: {e}")
            return

        chunks = self.chunk_by_sentences(full_text, chunk_size, chunk_overlap)
        documents_to_index = []
        for i, chunk in enumerate(chunks):
            doc_id = f"{source}-{i}-sentence"
            metadata = {"source": source}
            documents_to_index.append({"id": doc_id, "text": chunk, "metadata": metadata})
        self.add_documents(documents_to_index)
        print(f"Indexed {len(chunks)} sentence-based chunks from {file_path}")

    def retrieve(self, query: str, top_k: int = 5) -> List[Dict]:
        if not self.index:
            print("Pinecone index not initialized. Cannot retrieve.")
            return []

        query_embedding = self.embed_text(query)

        results = self.index.query(
            vector=query_embedding,
            top_k=top_k,
            include_metadata=True
        )

        retrieved_docs = []
        for match in results['matches']:
            retrieved_docs.append({
                'id': match['id'],
                'score': match['score'],
                'text': match['metadata'].get('text', ''),
                'metadata': match['metadata']
            })

        return retrieved_docs

    def query(self, question: str, top_k: int = 3) -> Dict:
        retrieved_docs = self.retrieve(question, top_k=top_k)
        context = "\n\n".join([doc['text'] for doc in retrieved_docs])
        answer = None
        if self.gemini_client:
            answer = self.generate(question, context)

        return {
            # "question": question,
            "answer": answer,
            # "retrieved_documents": retrieved_docs
        }

    def generate(self, query, context) -> str:
        if not self.gemini_client:
            raise ValueError("Gemini client not initialized. Please provide a Gemini API key.")

        prompt =  f"""
        <Profile>
        You are an efficient and effective RAG Chatbot, that helps students study and read information from their materials. You are designed to provide helpful and accurate responses to questions asked by students. 
        You are also designed to
        1. Interact with the user in a friendly and professional manner like a teacher. Use emojis and a friendly tone to engage with customers.
        
        <Responsibilities>
        NOTE:
        1. You do not give verbose and irrelevant responses instead, you give direct answers to questions asked. As part of your task,Use the following instructions below to answer the questions given to you at the end. Please follow the following rules:
        2. Use the context provided below.\
        3. If you don't know the answer or you can't find it in the context provided to you, don't try to make up an answer".\
        6. Choose the most appropriate answer among the list of responses from the vector search.\
        7. If you can't find the answer, don't try to make up an answer.\


        {context}

        Question: {query}

        Helpful Answer:
            """

        try:
            response = self.gemini_client.generate_content(prompt)
            return response.text
        except Exception as e:
            print(f"Error during Gemini generation: {e}")
            return "I encountered an error while trying to generate an answer."

    def split_text_into_sentences(self, text):
        sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s', text)
        return sentences
